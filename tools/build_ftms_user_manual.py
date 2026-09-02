from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "FTMS_User_Manual.docx"

GREEN = "0B5D45"
BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F5EF"
LIGHT_AMBER = "FFF4D6"
LIGHT_RED = "FCE8E8"
GRAY = "F2F4F7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_note(doc, title, text, fill=LIGHT_GREEN):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(GREEN)
    paragraph.add_run(f" {text}")
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="FTMS Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="FTMS Number")


def add_matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=BLUE)
        set_cell_shading(table.rows[0].cells[idx], LIGHT_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 24, GREEN, 0, 6),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    bullet = styles.add_style("FTMS Bullet", 1)
    bullet.base_style = normal
    bullet.paragraph_format.left_indent = Inches(0.375)
    bullet.paragraph_format.first_line_indent = Inches(-0.188)
    bullet.paragraph_format.space_after = Pt(4)

    number = styles.add_style("FTMS Number", 1)
    number.base_style = normal
    number.paragraph_format.left_indent = Inches(0.375)
    number.paragraph_format.first_line_indent = Inches(-0.188)
    number.paragraph_format.space_after = Pt(4)


def add_cover(doc):
    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FTMS User Manual")
    run.bold = True
    run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Foreign Travel Management System")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor.from_string(BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Prepared for Ministry of Agriculture | {date.today().strftime('%B %d, %Y')}")

    add_note(
        doc,
        "Manual purpose.",
        "This document explains the full FTMS operating process for travelers, approvers, protocol, PM Office follow-up, and administrators.",
    )
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Contents", level=1)
    add_numbered(
        doc,
        [
            "System overview and access model",
            "Roles and responsibilities",
            "Account registration, login, and password requests",
            "Organization setup and governance",
            "Creating, drafting, and submitting travel requests",
            "Approval workflow by traveler type",
            "Submitted requests and decision actions",
            "Travel status timeline, notifications, and reports",
            "Support letter and PM Office process",
            "Administration, audit, and troubleshooting",
        ],
    )
    doc.add_page_break()


def build_manual():
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_toc(doc)

    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph(
        "FTMS is the Ministry of Agriculture Foreign Travel Management System. It registers travelers, routes travel requests through the correct approval chain, records decisions, supports PM Office submission when required, and provides dashboards, reports, notifications, and audit history."
    )
    add_note(
        doc,
        "Core principle.",
        "The system routes the request according to the traveler's real structure and skips self-approval when the traveler is also an approver.",
        LIGHT_AMBER,
    )
    add_matrix(
        doc,
        ["Area", "What the user can do"],
        [
            ["Dashboard", "See counts, approval workload, historical totals, and current operating picture."],
            ["New Travel Request", "Create a new request, save a draft, complete a draft, and submit travel information."],
            ["Submitted Requests", "Review assigned requests, approve, reject, return/amend, or continue eligible drafts."],
            ["Travel Status", "View the request timeline with completed, current, and upcoming stages."],
            ["Reports", "Analyze approved travel by month, MoA vs Affiliate travel, sectors, organizations, and funding."],
            ["Administration", "Manage structures, projects, affiliate institutions, approvers, users, resets, and audit trail."],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("2. Roles and Responsibilities", level=1)
    add_matrix(
        doc,
        ["Role", "Main responsibility"],
        [
            ["Traveler / Expert", "Creates travel requests, saves drafts, submits completed requests, responds to returned requests, and tracks status."],
            ["Lead Executive Officer", "Reviews staff requests under the Lead Executive Office and forwards approved requests to the next approver."],
            ["Project Coordinator", "Reviews project staff requests, but never approves their own travel request."],
            ["Director General", "Default approver for travelers under an Affiliate Institute and may also submit personal travel requests."],
            ["State Minister", "Reviews sector-related requests and can submit personal travel requests without self-approval."],
            ["CEO", "Reviews CEO-structure requests and can submit personal travel requests without self-approval."],
            ["Head of the Minister's Office", "Makes final ministry office decision or forwards to the Minister for decision."],
            ["Protocol", "Reviews all travelers before final approval and selects whether PM Office approval is required."],
            ["Minister", "Decides requests forwarded by the Head of the Minister's Office and minister-structure requests."],
            ["PM Office", "Tracks PM Office submission and follow-up when PM approval is required."],
            ["Admin / Super Admin", "Maintains users, organizations, approvers, system settings, audit records, and reports."],
        ],
        [1.7, 4.8],
    )

    doc.add_heading("3. Account Registration and Login", level=1)
    add_numbered(
        doc,
        [
            "Open the FTMS web address and choose Login or Create Account.",
            "For a new traveler, enter personal information, email, organization type, structure, and role assignment.",
            "If the email already exists, the system asks the user to request password reset from the admin instead of creating another account.",
            "Admin reviews pending users and activates the account.",
            "After approval, the user logs in with the registered email and password.",
        ],
    )
    add_note(
        doc,
        "Duplicate account handling.",
        "When a traveler is already registered, the correct action is to request a password reset from the administrator. This keeps the user's request history linked to one account.",
    )

    doc.add_heading("4. Organization Setup", level=1)
    doc.add_paragraph("Organization Settings is the control point for workflow ownership. Admins should maintain it before users submit requests.")
    add_matrix(
        doc,
        ["Setup item", "Purpose"],
        [
            ["MoA Structures", "Registers Sector, CEO, Head of the Minister's Office, and Minister structures."],
            ["Lead Executive Offices", "Registers offices under sectors, CEO, or Head of the Minister's Office."],
            ["Projects", "Registers MoA projects under sectors, CEO, or Head of the Minister's Office."],
            ["Affiliate Institutions", "Registers external institutions and their Director General information."],
            ["Workflow Approvers", "Assigns State Ministers, CEO, Office Head, Minister, Protocol, Project Coordinators, and Lead Executives."],
            ["Hierarchy View", "Checks the relationship between structures, offices, projects, approvers, and organizations."],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("5. Creating a Travel Request", level=1)
    add_numbered(
        doc,
        [
            "Open New Travel Request.",
            "Select organization context: MoA or Affiliate Institute.",
            "Select your structure: Project Staff, Advisor, or Staff under Lead Executive.",
            "Select the matching project, advisor structure, affiliate institution, sector, CEO, office head, or lead executive office as required.",
            "Enter destination country. While typing, the country field filters the list.",
            "Enter travel dates, purpose, invitation details, funding source, and sponsor when funding is non-government.",
            "Upload or attach required supporting documents when requested by the system.",
            "Use Save as Draft when the request is incomplete; use Submit only when the travel request is ready for workflow routing.",
        ],
    )
    add_note(
        doc,
        "Draft behavior.",
        "A saved draft remains in Expert Preparation and appears in Submitted Requests with a Complete Draft action for the traveler only.",
        LIGHT_BLUE,
    )

    doc.add_heading("6. Workflow by Traveler Type", level=1)
    add_matrix(
        doc,
        ["Traveler type", "Default route"],
        [
            ["Staff under Lead Executive", "Expert Preparation -> Lead Executive Review -> State Minister/CEO/Office Head Review -> Protocol Clearance -> Office Head Final Decision -> PM Office if required -> Completed."],
            ["Advisor under Sector", "Expert Preparation -> State Minister Review -> Protocol Clearance -> Office Head Final Decision -> PM Office if required -> Completed."],
            ["Advisor under CEO", "Expert Preparation -> CEO Review -> Protocol Clearance -> Office Head Final Decision -> PM Office if required -> Completed."],
            ["Advisor under Minister or Office Head", "Expert Preparation -> Protocol Clearance -> Office Head Final Decision -> PM Office if required -> Completed."],
            ["Project Staff", "Expert Preparation -> Project Coordinator Review -> Parent Structure Approver -> Protocol Clearance -> Office Head Final Decision -> PM Office if required -> Completed."],
            ["Affiliate Institute Traveler", "Expert Preparation -> Director General Review -> Protocol Clearance -> Office Head Final Decision -> PM Office if required -> Completed."],
            ["Higher Official / Approver as Traveler", "The request skips that user's own approval stage and moves to the next valid approver."],
        ],
        [1.75, 4.75],
    )
    add_note(
        doc,
        "Self-approval protection.",
        "If the traveler is a Lead Executive, Project Coordinator, Director General, State Minister, CEO, Office Head, or Minister, FTMS treats the person as a traveler and passes the request to the next approver.",
        LIGHT_AMBER,
    )

    doc.add_heading("7. Protocol, Office Head, Minister, and PM Office", level=1)
    doc.add_paragraph("Protocol now reviews all travelers before final ministry decision. At Protocol Clearance, protocol selects one of two PM decision gates.")
    add_matrix(
        doc,
        ["Decision point", "Available action"],
        [
            ["Protocol Clearance", "Mark PM approval required or no PM approval required, then forward to Office Head Final Decision."],
            ["Office Head Final Decision", "Approve, Reject, or Forward to Minister with a note."],
            ["Minister Approval", "Approve or Reject a request forwarded by the Head of the Minister's Office or minister-structure request."],
            ["PM Office Submission", "Protocol submits approved requests that require PM Office clearance."],
            ["PM Office Follow-up", "PM Office user follows up and closes the PM approval stage."],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("8. Submitted Requests", level=1)
    add_bullets(
        doc,
        [
            "Approvers see requests assigned to their current role and stage.",
            "Travelers see their own draft, returned, pending, approved, rejected, and historical requests.",
            "Draft requests show Complete Draft for the traveler.",
            "Once decision review starts, the traveler cannot edit the request unless it is returned for amendment.",
            "On tablet or mobile web, request detail views should open in a popup-style layout to reduce scrolling.",
        ],
    )

    doc.add_heading("9. Travel Status Timeline", level=1)
    add_matrix(
        doc,
        ["Timeline state", "Color", "Information shown"],
        [
            ["Completed", "Green", "Approver name and approval/submission date."],
            ["Current", "Amber", "Current approver, date reached, and how many days the request has been pending in the current stage."],
            ["Upcoming", "Light red", "Upcoming approver name where the system can identify it."],
            ["Optional Minister path", "Dotted line", "Shows that Office Head may forward to Minister for decision."],
        ],
        [1.6, 1.2, 3.7],
    )

    doc.add_heading("10. Notifications and Reports", level=1)
    add_bullets(
        doc,
        [
            "Notification counts appear on relevant navigation items when a user has pending decisions or messages.",
            "Submitted Requests notifications update when requests are assigned, approved, rejected, returned, or completed.",
            "Reports include approved travel by month, total travel by MoA and Affiliate Institute, MoA travel by sector, Affiliate travel by organization, and funding-related summaries.",
            "Dashboards show total requests, approved, pending, rejected, and structure-based counts with compact interactive cards.",
        ],
    )

    doc.add_heading("11. Support Letter", level=1)
    doc.add_paragraph(
        "After the request is approved by the Office Head or Minister, FTMS generates the support letter used for PM Office submission when PM approval is required. The letter is addressed to ለጠቅላይ ሚኒስትር ጽ/ቤት and should not be addressed to the traveler's affiliate organization."
    )

    doc.add_heading("12. Administration and Audit", level=1)
    add_matrix(
        doc,
        ["Admin function", "Usage"],
        [
            ["Pending Users", "Approve new accounts after verifying their structure and role."],
            ["User Management", "Search and filter accounts by role, structure, organization, status, and responsibility group."],
            ["Organization Settings", "Maintain structures, lead offices, projects, affiliate institutions, and approvers."],
            ["Reset Password", "Support registered users who forgot credentials."],
            ["Audit Trail", "Review accountability records for approvals, edits, user changes, and workflow events."],
        ],
        [1.8, 4.7],
    )

    doc.add_heading("13. Troubleshooting", level=1)
    add_matrix(
        doc,
        ["Message or symptom", "Likely action"],
        [
            ["Unable to load countries or organization lists", "Check that the backend API is running and the frontend API URL points to the correct server."],
            ["Failed to load requests", "Confirm the backend is online, the route exists, and IIS/proxy is forwarding /api requests correctly."],
            ["No submitted requests visible", "Check the user's role, organization assignment, active status, and whether the request is at that user's workflow stage."],
            ["User exists during registration", "Use the password reset request flow instead of creating a duplicate account."],
            ["A draft has no workflow movement", "Open Submitted Requests and use Complete Draft, then submit the completed request."],
            ["Wrong approver shown", "Review Organization Settings hierarchy and the user's role/structure assignment."],
        ],
        [2.0, 4.5],
    )

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("FTMS User Manual | Ministry of Agriculture | Auto-generated documentation")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_manual()
