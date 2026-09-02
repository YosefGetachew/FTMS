from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/FTMS_Software_Requirements_Specification.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9E3DD")
        borders.append(tag)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        shade_cell(hdr[i], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_req_table(doc, rows):
    return add_table(
        doc,
        ["ID", "Requirement", "Priority", "Acceptance / Verification"],
        rows,
        widths=[0.75, 3.15, 0.8, 1.8],
    )


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "FTMS Software Requirements Specification"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Ministry of Agriculture - Foreign Travel Management System"
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Software Requirements Specification")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 81, 50)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Foreign Travel Management System (FTMS)")
    r.font.name = "Calibri"
    r.font.size = Pt(18)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Ministry of Agriculture")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    add_table(
        doc,
        ["Document Field", "Value"],
        [
            ["Document Type", "Software Requirements Specification (SRS)"],
            ["System", "Foreign Travel Management System"],
            ["Prepared For", "Ministry of Agriculture"],
            ["Prepared Date", "June 8, 2026"],
            ["Version", "1.0"],
            ["Prepared From", "Current FTMS source code, backend routes, web UI, mobile app, and deployment behavior"],
        ],
        widths=[1.8, 4.7],
    )
    doc.add_paragraph(
        "This SRS defines the functional, data, workflow, interface, security, reporting, mobile, and operational requirements for FTMS. It is intended to support development, testing, deployment, maintenance, and stakeholder review."
    )
    doc.add_page_break()


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_footer(doc)
    add_cover(doc)

    add_h1(doc, "Revision History")
    add_table(
        doc,
        ["Version", "Date", "Author", "Description"],
        [["1.0", "2026-06-08", "Codex / FTMS project workspace", "Initial intensive SRS generated from system implementation."]],
        widths=[0.8, 1.1, 2.1, 2.5],
    )

    add_h1(doc, "1. Introduction")
    add_h2(doc, "1.1 Purpose")
    doc.add_paragraph(
        "The purpose of this document is to specify the requirements for the Foreign Travel Management System (FTMS), a web and mobile-enabled system used to submit, route, approve, monitor, report, and audit foreign travel requests for Ministry of Agriculture personnel and affiliate institutions."
    )
    add_h2(doc, "1.2 Scope")
    doc.add_paragraph(
        "FTMS covers account registration and approval, role-based access, travel request creation, document attachment, workflow decisions, protocol clearance, Office Head and Minister review, PM Office follow-up, notifications, reporting, PDF generation, organization settings, user management, and audit trail visibility. The system includes a React web application, an Expo/React Native mobile application, an Express.js backend API, a PostgreSQL database, filesystem document storage, email delivery, and generated travel-request PDFs."
    )
    add_h2(doc, "1.3 Intended Audience")
    add_bullets(
        doc,
        [
            "Ministry leadership and business owners validating workflow coverage.",
            "System administrators configuring roles, users, structures, and affiliate institutions.",
            "Developers maintaining the frontend, backend, mobile app, API routes, and database.",
            "Testers preparing acceptance tests, regression tests, and deployment validation.",
            "Operations staff responsible for production deployment, backup, monitoring, and support.",
        ],
    )
    add_h2(doc, "1.4 Definitions and Acronyms")
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["FTMS", "Foreign Travel Management System."],
            ["MoA", "Ministry of Agriculture."],
            ["LEO", "Lead Executive Officer / Lead Executive Office, depending on context."],
            ["PM Office", "Prime Minister Office follow-up stage and role in the workflow."],
            ["RBAC", "Role-Based Access Control."],
            ["SRS", "Software Requirements Specification."],
        ],
        widths=[1.5, 5.0],
    )

    add_h1(doc, "2. Overall Description")
    add_h2(doc, "2.1 Product Perspective")
    doc.add_paragraph(
        "FTMS is a custom business workflow application. It centralizes foreign travel request data and replaces manual or fragmented approval tracking with a role-based electronic workflow. The backend exposes REST-style JSON and multipart endpoints, the web frontend provides the full administrative and operational surface, and the mobile app provides role-aware access to core workflow and reporting features."
    )
    add_h2(doc, "2.2 Product Functions")
    add_bullets(
        doc,
        [
            "User registration, login, JWT session generation, and pending account approval.",
            "Travel request creation for MoA and affiliate institution travelers.",
            "Upload of passport copy, invitation letter, and Terms of Reference documents.",
            "Stage-based workflow routing across Expert, Lead Executive, State Minister, CEO, Office Head, Protocol, Minister, and PM Office roles.",
            "Email and in-app notifications for account and travel workflow events.",
            "Dashboard statistics, reports, monthly trends, sector summaries, stage summaries, and Office Head/Minister reports.",
            "PDF generation for travel request summaries, including English and Amharic date context.",
            "Organization settings for MoA sectors/offices, executive offices, affiliate institutions, and approver setup.",
            "Audit trail of workflow actions, actors, comments, status transitions, and stage transitions.",
        ],
    )
    add_h2(doc, "2.3 User Classes and Characteristics")
    add_table(
        doc,
        ["User Class", "Primary Responsibilities", "Expected Skill Level"],
        [
            ["Traveler / Expert", "Register, create travel requests, attach documents, submit or resubmit corrected requests, track status.", "Basic web/mobile user."],
            ["Lead Executive Officer", "Review assigned requests from the same sector and department.", "Operational approver."],
            ["State Minister", "Review sector-structure requests assigned to their sector.", "Executive approver."],
            ["CEO / Chief Executive Officer", "Review CEO-structure requests.", "Executive approver."],
            ["Office Head", "Review office-head-structure requests, final office review, reports.", "Executive/administrative approver."],
            ["Protocol", "Perform protocol clearance, request amendment, submit requests to PM Office, view audit trail.", "Workflow operator."],
            ["Minister", "Perform final approval stage and view ministry-level reports.", "Executive approver."],
            ["PM Office", "Approve or reject PM Office follow-up stage.", "External/coordination approver."],
            ["Admin / Super Admin", "Manage users, pending users, organization settings, reports, audit, and all workflow visibility.", "System administrator."],
        ],
        widths=[1.45, 3.65, 1.4],
    )
    add_h2(doc, "2.4 Operating Environment")
    add_table(
        doc,
        ["Layer", "Technology / Environment"],
        [
            ["Web Client", "React, React DOM, Axios, Recharts, Socket.IO client, browser-based deployment."],
            ["Mobile Client", "Expo SDK 54, React Native 0.81, Expo Router, fetch-based API client, iOS testing via Expo Go."],
            ["Backend", "Node.js, Express.js, PostgreSQL pg driver, bcryptjs, JSON Web Tokens, Multer, Nodemailer, PDFKit."],
            ["Database", "PostgreSQL database containing users, requests, audit trails, notifications, MoA settings, and affiliate institutions."],
            ["File Storage", "Backend filesystem uploads folder for PDF/JPG/PNG attachments and archive folder for completed requests."],
            ["Email", "SMTP transport configured by environment variables and Gmail SMTP host in current implementation."],
            ["Production Hosting", "Windows/IIS/PM2 deployment pattern observed during production troubleshooting."],
        ],
        widths=[1.4, 5.1],
    )
    add_h2(doc, "2.5 Assumptions and Constraints")
    add_bullets(
        doc,
        [
            "The backend must have reliable connectivity to the production PostgreSQL database.",
            "Email delivery requires valid EMAIL_USER and EMAIL_PASS configuration.",
            "Uploaded files are limited to PDF, JPG, and PNG formats and 5 MB per file.",
            "Mobile production distribution to iOS users requires Apple Developer credentials and EAS build configuration.",
            "Some backend routes currently allow query-based role access; hardening requirements are specified in security requirements.",
            "The system language is primarily English, with generated PDFs supporting Amharic date output where configured.",
        ],
    )

    add_h1(doc, "3. System Architecture")
    add_h2(doc, "3.1 Architectural Overview")
    add_table(
        doc,
        ["Component", "Responsibilities", "Interfaces"],
        [
            ["React Web Frontend", "Full-featured user interface for workflow, admin, reporting, settings, notifications, and audit.", "REST API over HTTPS."],
            ["Expo Mobile App", "Role-aware mobile access to dashboard, requests, reports, admin screens, audit, and password change.", "REST API over HTTPS."],
            ["Express Backend", "Authentication, workflow rules, file upload, notifications, reporting, PDF generation, settings, users.", "HTTP JSON/multipart API, SMTP, PostgreSQL."],
            ["PostgreSQL Database", "Persistent storage for users, requests, workflow stages, audit records, notifications, and settings.", "SQL via pg Pool."],
            ["Filesystem Storage", "Stores uploaded travel documents and generated PDFs.", "Backend local paths served through Express."],
            ["SMTP Service", "Sends account and workflow emails using configured sender name MoA-Foreign Travel.", "Nodemailer SMTP transport."],
        ],
        widths=[1.45, 3.45, 1.6],
    )
    add_h2(doc, "3.2 Logical Data Flow")
    add_numbered(
        doc,
        [
            "A user registers or is created by an administrator.",
            "Admin/Super Admin approves pending accounts where required.",
            "User logs in and receives a JWT and user profile.",
            "Traveler creates a request and uploads supporting documents.",
            "Traveler submits the request into the configured workflow path.",
            "Assigned approvers approve, reject, amend, clear, forward, submit to PM Office, or finalize the request.",
            "The backend records audit entries, updates statuses, sends notifications, and sends workflow emails.",
            "Leadership and administrators view dashboards, reports, PDFs, and audit trails.",
        ],
    )

    add_h1(doc, "4. Workflow Model")
    add_h2(doc, "4.1 Workflow Stages")
    add_table(
        doc,
        ["Stage Code", "Display Name", "Responsible Role"],
        [
            ["expert_preparation", "Expert Preparation", "traveler / expert"],
            ["lead_executive_review", "Lead Executive Officer Review", "lead_executive_officer / lead_executive"],
            ["state_minister_review", "State Minister Review", "state_minister"],
            ["ceo_review", "CEO Review", "chief_executive_officer / ceo"],
            ["office_head_review", "Office Head Review", "office_head"],
            ["protocol_clearance", "Protocol Clearance", "protocol"],
            ["office_head_final", "Office Head Final Decision", "office_head"],
            ["minister_review", "Minister Approval", "minister"],
            ["pm_office_submission", "Protocol Submission to PM Office", "protocol"],
            ["pm_office_followup", "PM Office Follow-up", "pm_office"],
            ["completed", "Completed", "System final state"],
        ],
        widths=[1.65, 2.8, 2.05],
    )
    add_h2(doc, "4.2 Workflow Paths")
    add_table(
        doc,
        ["Workflow Type", "Path"],
        [
            ["sector_structure", "Expert -> Lead Executive Officer -> State Minister -> Protocol -> Office Head Final -> Minister -> Protocol PM Submission -> PM Office -> Completed"],
            ["ceo_structure", "Expert -> Lead Executive Officer -> CEO -> Protocol -> Office Head Final -> Minister -> Protocol PM Submission -> PM Office -> Completed"],
            ["office_head_structure", "Expert -> Lead Executive Officer -> Office Head -> Protocol -> Office Head Final -> Minister -> Protocol PM Submission -> PM Office -> Completed"],
            ["affiliate_institution", "Expert -> Protocol -> Office Head Final -> Minister -> Protocol PM Submission -> PM Office -> Completed"],
        ],
        widths=[1.55, 4.95],
    )
    add_h2(doc, "4.3 Workflow Actions")
    add_table(
        doc,
        ["Action", "Valid Stage", "Result"],
        [
            ["submit", "expert_preparation", "Moves request to Lead Executive Review or Protocol Clearance for affiliate requests."],
            ["approve", "lead_executive_review, state_minister_review, ceo_review, office_head_review, minister_review", "Moves request to the next configured stage."],
            ["reject", "review stages", "Returns for amendment or final rejection depending on stage."],
            ["clear", "protocol_clearance", "Moves request to Office Head Final."],
            ["amend", "protocol_clearance", "Returns request to Expert Preparation with amendment comment."],
            ["forward_to_minister", "office_head_final", "Moves request to Minister Review."],
            ["submit_to_pm_office", "pm_office_submission", "Moves request to PM Office Follow-up."],
            ["pm_office_approved", "pm_office_followup", "Finalizes request as approved."],
            ["pm_office_rejected", "pm_office_followup", "Finalizes request as rejected."],
            ["resubmit", "expert_preparation with amended final status", "Resubmits corrected request to the appropriate next stage."],
        ],
        widths=[1.45, 2.2, 2.85],
    )

    add_h1(doc, "5. Functional Requirements")
    add_h2(doc, "5.1 Authentication and Account Lifecycle")
    add_req_table(
        doc,
        [
            ["FR-AUTH-001", "The system shall allow users to register as travelers with full name, email, password, phone, position, organization type, organization name, sector, and department.", "High", "Registration persists a pending user and prevents duplicate email addresses."],
            ["FR-AUTH-002", "The system shall require Admin or Super Admin approval before pending registered accounts can log in.", "High", "Pending accounts receive a 403 response until approved."],
            ["FR-AUTH-003", "The system shall authenticate users by normalized email and password.", "High", "Valid credentials return a JWT token and user profile; invalid credentials return an error."],
            ["FR-AUTH-004", "The system shall allow eligible users to change their password after validating the current password and new password strength.", "High", "Password change endpoint rejects missing, weak, or incorrect current passwords."],
            ["FR-AUTH-005", "The system shall allow administrators to create, update, disable, delete, and reset passwords for non-super-admin users.", "High", "Admin user management screens and API operations update user records."],
        ]
    )
    add_h2(doc, "5.2 Role-Based Access Control")
    add_req_table(
        doc,
        [
            ["FR-RBAC-001", "The system shall restrict visible requests by user role, email, sector, department, workflow type, and audit history as applicable.", "High", "Each role query returns only assigned, owned, historical, or administrative records."],
            ["FR-RBAC-002", "The system shall prevent approvers from deciding requests outside their assigned role and stage.", "High", "Status update route returns an error when role stage assignment does not match."],
            ["FR-RBAC-003", "The system shall restrict pending-user approval to Admin and Super Admin users.", "High", "Protected endpoint requires a valid JWT and admin role."],
            ["FR-RBAC-004", "The web and mobile interfaces shall hide or disable features unavailable to the logged-in role.", "Medium", "Minister and PM Office do not see traveler-only controls; admin pages require admin role."],
        ]
    )
    add_h2(doc, "5.3 Travel Request Management")
    add_req_table(
        doc,
        [
            ["FR-TRV-001", "The system shall allow authorized users to create MoA or affiliate travel requests.", "High", "Request is created with traveler, organization, travel dates, purpose, sponsor, passport number, contact, and workflow fields."],
            ["FR-TRV-002", "The system shall validate required request fields: full name, email, country, start date, and end date.", "High", "Request creation fails with validation error if required fields are missing."],
            ["FR-TRV-003", "The system shall support upload of passport file, invitation letter, and TOR file.", "High", "Only PDF, JPG, PNG files under 5 MB are accepted."],
            ["FR-TRV-004", "The system shall allow permitted users to update existing request details and attachments.", "Medium", "PUT request updates provided fields while preserving unspecified values."],
            ["FR-TRV-005", "The system shall allow request deletion through backend API for authorized operational use.", "Medium", "Deleting a request removes it and cascades related audit rows where configured."],
        ]
    )
    add_h2(doc, "5.4 Workflow Decisions")
    add_req_table(
        doc,
        [
            ["FR-WF-001", "The system shall route submitted requests according to workflow type and traveler category.", "High", "Sector, CEO, Office Head, and affiliate paths follow the stage definitions in Section 4."],
            ["FR-WF-002", "The system shall record every workflow action in request_audit_trails.", "High", "Audit entry includes request ID, action, actor role, actor email, comment, old/new stage, old/new status."],
            ["FR-WF-003", "The system shall notify the traveler when a request is updated, returned, approved, or rejected.", "High", "Notification record and email are created where configured."],
            ["FR-WF-004", "The system shall archive uploaded documents when a request reaches a final approved or rejected state.", "Medium", "Completed request files are moved to uploads/archive."],
            ["FR-WF-005", "The system shall support bulk Protocol submission to PM Office for requests ready at pm_office_submission.", "Medium", "Bulk endpoint updates eligible selected requests and skips ineligible ones."],
        ]
    )
    add_h2(doc, "5.5 Notifications and Email")
    add_req_table(
        doc,
        [
            ["FR-NOT-001", "The system shall create in-app notifications for travel request creation and workflow updates.", "High", "Notifications are visible by user email and ordered by date."],
            ["FR-NOT-002", "The system shall allow a user to mark one notification or all notifications as read.", "Medium", "Read status changes in notifications table."],
            ["FR-NOT-003", "The system shall send email notifications for account approval/rejection and workflow tasks where recipient email is available.", "High", "Email sender display name is MoA-Foreign Travel."],
            ["FR-NOT-004", "Email failures shall not block the main workflow transaction.", "Medium", "sendEmailSafe logs email errors and allows workflow operation to continue."],
        ]
    )
    add_h2(doc, "5.6 Reporting and Dashboard")
    add_req_table(
        doc,
        [
            ["FR-REP-001", "The system shall display request totals, approved, pending, rejected, and returned/amended counts.", "High", "Dashboard/stat endpoints calculate counts from requests."],
            ["FR-REP-002", "The system shall provide status summary, sector status, monthly approved requests, stage summary, and Office Head/Minister reports.", "High", "Reports endpoints return JSON suitable for charts and tables."],
            ["FR-REP-003", "The web interface shall visualize reports using charts and summary cards.", "Medium", "Reports component uses Recharts for line/bar/pie visualizations."],
            ["FR-REP-004", "The mobile interface shall expose reports to authorized roles and display mobile-friendly line graphs and summaries.", "Medium", "Reports tab visible to Admin, Super Admin, Minister, and Office Head."],
        ]
    )
    add_h2(doc, "5.7 Administration and Settings")
    add_req_table(
        doc,
        [
            ["FR-ADM-001", "The system shall allow administrators to manage MoA sectors/offices and their workflow type.", "High", "CRUD routes exist for /api/moa-sectors."],
            ["FR-ADM-002", "The system shall allow administrators to manage MoA executive offices under sectors/offices.", "High", "CRUD routes exist for /api/moa-executive-offices."],
            ["FR-ADM-003", "The system shall allow administrators to manage affiliate institutions.", "High", "CRUD routes exist for /api/affiliate-institutions."],
            ["FR-ADM-004", "The system shall allow administrators to configure workflow approver accounts.", "Medium", "State-ministers/approvers routes create, update, list, and delete workflow users."],
            ["FR-ADM-005", "The system shall seed a super admin account when no super admin exists.", "Medium", "Startup seed checks for super_admin and inserts default where needed."],
        ]
    )
    add_h2(doc, "5.8 PDF Generation")
    add_req_table(
        doc,
        [
            ["FR-PDF-001", "The system shall generate a PDF summary for a travel request by ID.", "High", "GET /api/generate-pdf/:id returns a downloadable PDF."],
            ["FR-PDF-002", "The PDF shall include traveler, organizational, workflow, destination, date, purpose, sponsor, passport, status, and signature fields.", "High", "Generated document contains English and Amharic summary pages."],
            ["FR-PDF-003", "The PDF generator shall use available Ethiopic font support when present.", "Medium", "NotoSansEthiopic-Regular.ttf is used if found."],
        ]
    )
    add_h2(doc, "5.9 Mobile Application")
    add_req_table(
        doc,
        [
            ["FR-MOB-001", "The mobile app shall support login, registration, dashboard, requests, status, notifications, profile, reports, More menu, admin screens, audit, settings, and password change by role.", "High", "Expo app exposes role-aware tabs and stack screens."],
            ["FR-MOB-002", "The mobile app shall connect to the production API using configurable EXPO_PUBLIC_API_URL.", "High", "API client defaults to https://ftms.moa.gov.et/api when env value is absent."],
            ["FR-MOB-003", "The mobile app shall send Authorization bearer token for authenticated API calls.", "High", "API client sets Authorization header after sign-in."],
            ["FR-MOB-004", "The mobile app shall support file selection for travel request attachments.", "Medium", "DocumentPicker accepts PDF/JPG/PNG and posts FormData."],
        ]
    )

    add_h1(doc, "6. External Interface Requirements")
    add_h2(doc, "6.1 User Interfaces")
    add_table(
        doc,
        ["Interface", "Required Capabilities"],
        [
            ["Web login/register", "Authenticate existing users and submit account registration requests."],
            ["Web dashboard", "Role-aware statistics, workload overview, and navigation to reports or requests."],
            ["Request form/table/status", "Create, review, decide, track, amend, resubmit, and inspect requests."],
            ["Reports", "Render status, sector, monthly, workflow stage, and organization-level analytics."],
            ["Settings", "Manage MoA structures, executive offices, affiliate institutions, and approvers."],
            ["Mobile app", "Provide role-aware access to core and administrative features optimized for iPhone screens."],
        ],
        widths=[1.7, 4.8],
    )
    add_h2(doc, "6.2 API Interfaces")
    add_table(
        doc,
        ["API Group", "Representative Routes"],
        [
            ["Auth", "POST /api/register, POST /api/login"],
            ["Requests", "POST /api/requests, GET /api/requests, PUT /api/requests/:id, PUT /api/requests/:id/status, PUT /api/requests/:id/resubmit, DELETE /api/requests/:id"],
            ["Audit", "GET /api/audit-trail, GET /api/requests/:id/audit-trail"],
            ["Users", "GET /api/users, GET /api/users/pending, POST /api/users, PUT /api/users/:id, approve/reject/reset/change-password routes"],
            ["Reports", "GET /api/stats, /api/dashboard/pending-by-sector, /api/reports/status-summary, sector-status, monthly-requests, stage-summary, office-minister-summary"],
            ["Settings", "CRUD /api/moa-sectors, /api/moa-executive-offices, /api/affiliate-institutions, /api/state-ministers"],
            ["Notifications", "GET /api/notifications, PUT /api/notifications/:id/read, PUT /api/notifications/read-all/:email"],
            ["PDF", "GET /api/generate-pdf/:id"],
        ],
        widths=[1.4, 5.1],
    )
    add_h2(doc, "6.3 Hardware and Software Interfaces")
    add_bullets(
        doc,
        [
            "Client browsers shall support modern JavaScript, React rendering, HTTPS, and file upload.",
            "iOS mobile testing shall run in Expo Go for development; production distribution shall use EAS and Apple credentials.",
            "Backend server shall run Node.js and have write access to uploads and pdfs directories.",
            "PostgreSQL shall be reachable from the backend using DB_USER, DB_HOST, DB_NAME, DB_PASSWORD, and DB_PORT.",
            "SMTP service shall be reachable using EMAIL_USER and EMAIL_PASS.",
        ],
    )

    add_h1(doc, "7. Data Requirements")
    add_h2(doc, "7.1 Core Entities")
    add_table(
        doc,
        ["Entity / Table", "Purpose", "Representative Fields"],
        [
            ["users", "Stores account, role, organization, and approval state.", "id, full_name, email, password, role, phone, position, sector, department, account_status, is_active"],
            ["requests", "Stores current travel request records and workflow state.", "id, traveler_category, workflow_type, current_stage, final_status, full_name, sector, country, start_date, end_date, attachments, status"],
            ["request_audit_trails", "Stores workflow decision history.", "request_id, action, actor_role, actor_email, comment, old_stage, new_stage, old_status, new_status, created_at"],
            ["notifications", "Stores in-app notification messages.", "user_email, title, message, is_read, created_at"],
            ["moa_sectors", "Stores MoA sector/office settings.", "id, name, workflow_type, created_at"],
            ["moa_executive_offices", "Stores executive offices under sectors/offices.", "id, sector_id, name/office_name, created_at"],
            ["affiliate_institutions", "Stores affiliate institutions and contact info.", "id, organization_name, general_director_name, email, phone, created_at"],
            ["travel_requests", "Legacy travel request table observed in production backups.", "Legacy data retained or migrated separately."],
        ],
        widths=[1.55, 2.25, 2.7],
    )
    add_h2(doc, "7.2 Data Validation Requirements")
    add_bullets(
        doc,
        [
            "Email addresses shall be normalized by trimming and converting to lowercase before uniqueness or lookup checks.",
            "Passwords shall be hashed using bcrypt before storage.",
            "Phone numbers shall be normalized to Ethiopian +251 format where possible.",
            "Travel dates shall support ISO-style parsing for date display, duration, and reporting.",
            "Uploaded filenames shall be sanitized and prefixed with a timestamp.",
            "Request stage and status values shall be constrained to recognized workflow values.",
        ],
    )
    add_h2(doc, "7.3 Data Retention and Archiving")
    add_bullets(
        doc,
        [
            "Audit trail records should be retained for accountability and not deleted during normal request lifecycle operations.",
            "Uploaded documents should be retained while a request is active and moved to archive when finalized.",
            "Generated PDFs may be regenerated from request data and should not be treated as the source of truth.",
            "Production database backups should be taken before destructive maintenance, schema migration, or bulk deletion.",
        ],
    )

    add_h1(doc, "8. Non-Functional Requirements")
    add_h2(doc, "8.1 Security")
    add_req_table(
        doc,
        [
            ["NFR-SEC-001", "The system shall hash all stored passwords using bcrypt.", "High", "No plaintext passwords are stored."],
            ["NFR-SEC-002", "The system shall issue signed JWTs for authenticated sessions.", "High", "Token contains user ID, email, and role and expires after configured duration."],
            ["NFR-SEC-003", "Sensitive configuration shall be stored in environment variables rather than committed source code.", "High", "DB and email credentials read from .env."],
            ["NFR-SEC-004", "Administrative and decision APIs shall enforce authentication and role authorization.", "High", "Protected endpoints reject missing/invalid tokens and non-admin roles where applicable."],
            ["NFR-SEC-005", "File uploads shall restrict file type and file size.", "High", "Multer filter allows only PDF/JPG/PNG and 5 MB limit."],
            ["NFR-SEC-006", "Future hardening shall remove trust in client-supplied role values for decision authorization.", "High", "Backend derives actor role from validated JWT or server-side user lookup."],
        ]
    )
    add_h2(doc, "8.2 Performance")
    add_req_table(
        doc,
        [
            ["NFR-PERF-001", "Dashboard and report endpoints should return within 3 seconds for normal production data volumes.", "Medium", "Report queries execute within target under expected load."],
            ["NFR-PERF-002", "File upload requests should complete within 10 seconds for files under the configured 5 MB limit on normal networks.", "Medium", "Upload operation finishes or returns a clear error."],
            ["NFR-PERF-003", "Mobile screens should render without blocking on unused administrative data for unauthorized roles.", "Medium", "Role-aware navigation avoids unnecessary calls."],
        ]
    )
    add_h2(doc, "8.3 Availability and Reliability")
    add_req_table(
        doc,
        [
            ["NFR-REL-001", "The backend process shall be restartable by production process manager.", "High", "PM2 restart restores API service."],
            ["NFR-REL-002", "Email delivery errors shall be logged without rolling back core workflow updates.", "Medium", "Workflow continues when SMTP fails."],
            ["NFR-REL-003", "The application shall provide clear API errors for missing routes, validation failures, and unauthorized operations.", "Medium", "JSON error payload includes error message."],
            ["NFR-REL-004", "Database backup and restore procedures shall be documented for production maintenance.", "High", "Operators can run pg_dump/restore or SQL backup before migrations."],
        ]
    )
    add_h2(doc, "8.4 Usability")
    add_req_table(
        doc,
        [
            ["NFR-USE-001", "The web interface shall present role-appropriate navigation groups.", "Medium", "Sidebar sections show only permitted workspaces."],
            ["NFR-USE-002", "The mobile interface shall present all available system features in a role-aware More menu.", "Medium", "More screen groups features by category and role."],
            ["NFR-USE-003", "Workflow buttons shall use labels that match the current stage action.", "Medium", "Buttons show Approve, Reject, Clear, Amend, Submit PM, PM OK, etc."],
            ["NFR-USE-004", "Reports shall use charts and summaries appropriate to the device form factor.", "Medium", "Web uses Recharts; mobile uses compact summaries and line graphs."],
        ]
    )
    add_h2(doc, "8.5 Compatibility")
    add_req_table(
        doc,
        [
            ["NFR-COMP-001", "The web frontend shall support modern desktop browsers used by Ministry staff.", "Medium", "Production build runs through react-scripts build."],
            ["NFR-COMP-002", "The mobile app shall be compatible with Expo SDK 54 and iOS Expo Go during testing.", "High", "Expo export for iOS succeeds."],
            ["NFR-COMP-003", "The backend shall support Windows production deployment with PM2 and IIS reverse proxy patterns.", "Medium", "API responds through configured public domain."],
        ]
    )

    add_h1(doc, "9. Business Rules")
    add_table(
        doc,
        ["Rule ID", "Business Rule"],
        [
            ["BR-001", "Only active, approved accounts may log in."],
            ["BR-002", "A request cannot be updated through workflow once it is completed, approved, or rejected."],
            ["BR-003", "A role may decide only requests currently assigned to one of that role's stages."],
            ["BR-004", "Lead Executive review shall be scoped to matching sector and department."],
            ["BR-005", "State Minister review shall apply only to sector_structure workflow and matching sector."],
            ["BR-006", "CEO review shall apply only to ceo_structure workflow and matching CEO structure."],
            ["BR-007", "Protocol may clear or amend only at protocol_clearance and submit to PM Office only at pm_office_submission."],
            ["BR-008", "PM Office approval or rejection finalizes the request."],
            ["BR-009", "Affiliate institution requests enter the workflow at Protocol Clearance after submission."],
            ["BR-010", "Admin and Super Admin have administrative visibility over all non-final and reporting workflows."],
        ],
        widths=[0.85, 5.65],
    )

    add_h1(doc, "10. Acceptance Test Scenarios")
    add_table(
        doc,
        ["Scenario ID", "Scenario", "Expected Result"],
        [
            ["AT-001", "Register a new MoA traveler account.", "Account is created as pending and cannot log in until approved."],
            ["AT-002", "Admin approves a pending account.", "Account becomes active, user can log in, and approval email is attempted."],
            ["AT-003", "Traveler creates and submits a sector_structure request.", "Request moves to lead_executive_review and audit/notification records are created."],
            ["AT-004", "Lead Executive approves sector_structure request.", "Request moves to state_minister_review."],
            ["AT-005", "State Minister approves request.", "Request moves to protocol_clearance."],
            ["AT-006", "Protocol requests amendment.", "Request returns to expert_preparation with amended final status and comment."],
            ["AT-007", "Traveler resubmits amended request.", "Request moves back to the correct next stage."],
            ["AT-008", "Office Head forwards final request.", "Request moves to minister_review."],
            ["AT-009", "Minister approves request.", "Request moves to pm_office_submission for Protocol."],
            ["AT-010", "Protocol submits to PM Office.", "Request moves to pm_office_followup and PM Office status becomes submitted."],
            ["AT-011", "PM Office approves request.", "Request final_status becomes approved, current_stage becomes completed, files are archived."],
            ["AT-012", "Minister opens reports on mobile.", "Reports tab is visible and displays line graphs and summaries."],
            ["AT-013", "Admin opens organization settings.", "MoA sectors, executive offices, and affiliate institutions can be created, listed, and deleted."],
            ["AT-014", "Protocol opens audit trail.", "Audit records are visible and searchable."],
        ],
        widths=[0.85, 3.35, 2.3],
    )

    add_h1(doc, "11. Traceability Matrix")
    add_table(
        doc,
        ["Requirement Area", "Backend Route / Module", "Web Component", "Mobile Screen"],
        [
            ["Authentication", "/api/register, /api/login, /api/users/:id/change-password", "Login, Register, ResetPassword", "login, register, account/password"],
            ["Request creation", "POST /api/requests, middleware/upload.js", "RequestForm", "request/new"],
            ["Request review", "GET /api/requests, PUT /api/requests/:id/status", "RequestTable", "(tabs)/requests"],
            ["Travel status", "GET /api/requests", "TravelStatus", "(tabs)/status"],
            ["Notifications", "/api/notifications", "Notifications", "(tabs)/notifications"],
            ["Reports", "/api/reports/*, /api/stats", "Reports, DashboardStats", "(tabs)/reports, dashboard"],
            ["Pending users", "/api/users/pending, approve, reject", "PendingUsers", "admin/pending-users"],
            ["User management", "/api/users CRUD/reset", "UserManagement", "admin/users"],
            ["Settings", "/api/moa-sectors, moa-executive-offices, affiliate-institutions", "Settings", "admin/settings"],
            ["Audit", "/api/audit-trail, /api/requests/:id/audit-trail", "AuditTrail", "admin/audit"],
            ["PDF", "/api/generate-pdf/:id", "Request details/actions", "request/[id]"],
        ],
        widths=[1.35, 2.2, 1.55, 1.4],
    )

    add_h1(doc, "12. Risks, Gaps, and Recommendations")
    add_table(
        doc,
        ["Risk / Gap", "Impact", "Recommendation"],
        [
            ["Client-supplied role in some workflow calls", "A malicious client could attempt unauthorized decisions if backend hardening is incomplete.", "Derive role and actor identity from JWT on every protected workflow route."],
            ["Mixed protected and unprotected routes", "Administrative data may be exposed if network access is not otherwise controlled.", "Apply authenticateUser and role middleware consistently to sensitive routes."],
            ["Filesystem-based uploads", "Files may be lost or hard to scale across servers.", "Use backup policy or object storage for production attachments."],
            ["Email provider dependency", "SMTP outages may delay user awareness.", "Monitor email failures and provide retry/queue capability."],
            ["Manual production database operations", "Wrong database or table deletion can cause data loss or failed maintenance.", "Document production DB host/name and require backup before destructive actions."],
            ["Mobile App Store distribution not configured", "iOS users cannot install production app outside Expo Go/TestFlight.", "Obtain Apple Developer account and configure EAS build/submission."],
        ],
        widths=[2.0, 2.1, 2.4],
    )

    add_h1(doc, "13. Appendix A - API Endpoint Catalog")
    endpoint_rows = [
        ["GET", "/", "Backend health message."],
        ["GET", "/api/health", "API health check."],
        ["POST", "/api/register", "Create pending traveler account."],
        ["POST", "/api/login", "Authenticate and return JWT/user profile."],
        ["POST", "/api/requests", "Create travel request with optional files."],
        ["GET", "/api/requests", "List requests by role/email/id query context."],
        ["PUT", "/api/requests/:id", "Update request details and attachments."],
        ["PUT", "/api/requests/:id/status", "Apply workflow action."],
        ["PUT", "/api/requests/bulk/submit-to-pm-office", "Bulk submit eligible requests to PM Office."],
        ["PUT", "/api/requests/:id/resubmit", "Resubmit amended request."],
        ["DELETE", "/api/requests/:id", "Delete request."],
        ["GET", "/api/audit-trail", "List global audit records."],
        ["GET", "/api/requests/:id/audit-trail", "List audit records for one request."],
        ["GET", "/api/users", "List users."],
        ["GET", "/api/users/pending", "List pending users."],
        ["POST", "/api/users", "Create user."],
        ["PUT", "/api/users/:id", "Update user."],
        ["PUT", "/api/users/:id/approve", "Approve pending user."],
        ["PUT", "/api/users/:id/reject", "Reject pending user."],
        ["DELETE", "/api/users/:id", "Delete user."],
        ["PUT", "/api/users/:id/reset-password", "Admin reset user password."],
        ["PUT", "/api/users/:id/change-password", "User password change."],
        ["GET", "/api/stats", "Dashboard counts."],
        ["GET", "/api/dashboard/pending-by-sector", "Pending workload by sector."],
        ["GET", "/api/reports/status-summary", "Report counts by final status."],
        ["GET", "/api/reports/sector-status", "Report counts by sector/status."],
        ["GET", "/api/reports/monthly-requests", "Approved requests grouped by month."],
        ["GET", "/api/reports/stage-summary", "Requests grouped by current stage."],
        ["GET", "/api/reports/office-minister-summary", "MoA/affiliate and organization summaries."],
        ["GET", "/api/generate-pdf/:id", "Generate request PDF."],
        ["GET/POST/PUT/DELETE", "/api/moa-sectors", "Manage MoA sector/office settings."],
        ["GET/POST/PUT/DELETE", "/api/moa-executive-offices", "Manage executive office settings."],
        ["GET/POST/PUT/DELETE", "/api/affiliate-institutions", "Manage affiliate institutions."],
        ["GET/POST/PUT/DELETE", "/api/state-ministers", "Manage workflow approvers."],
        ["GET", "/api/notifications", "List notifications by email."],
        ["PUT", "/api/notifications/:id/read", "Mark one notification read."],
        ["PUT", "/api/notifications/read-all/:email", "Mark all notifications read."],
    ]
    add_table(doc, ["Method", "Endpoint", "Purpose"], endpoint_rows, widths=[1.0, 2.35, 3.15])

    add_h1(doc, "14. Appendix B - Deployment and Configuration Requirements")
    add_table(
        doc,
        ["Configuration", "Description"],
        [
            ["DB_USER, DB_HOST, DB_NAME, DB_PASSWORD, DB_PORT", "PostgreSQL connection variables used by backend/config/db.js."],
            ["JWT_SECRET", "Token signing secret. Must be strong and environment-specific."],
            ["EMAIL_USER, EMAIL_PASS", "SMTP credentials used by Nodemailer."],
            ["FRONTEND_URL", "Allowed frontend URL in deployment configuration."],
            ["EXPO_PUBLIC_API_URL", "Mobile app API base URL, normally https://ftms.moa.gov.et/api in production."],
            ["PORT", "Backend listening port, default 5000."],
        ],
        widths=[2.1, 4.4],
    )
    doc.add_paragraph(
        "Operational note: production maintenance tasks such as database dump, restore, schema migration, and bulk deletion must begin with a verified backup and must confirm the backend is connected to the intended database."
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
